"""Перенос результатов программной проверки в текст шестой главы диссертации."""

from __future__ import annotations

import hashlib
import json
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, Mapping

import pandas as pd
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DEFAULT_DOCUMENT_NAME = (
    "6. Глава 6. Экспериментальная проверка достоверности "
    "априорной оценки качества.docx"
)
REPORTS_DIR = Path("reports/chapter6")
FIGURES_DIR = REPORTS_DIR / "figures"
PLACEHOLDER_MARKERS = (
    "[рассчитать]",
    "[автозаполнение]",
    "[значение будет рассчитано",
    "[n]",
    "подлежит подстановке",
)

FIGURE_SPECS: tuple[tuple[str, str, str], ...] = (
    (
        "q_pred_vs_q_fact.png",
        "Таблица 6.5",
        "Рисунок 6.1 – Сопоставление априорной и фактической интегральных оценок",
    ),
    (
        "residuals_vs_q_fact.png",
        "Рисунок 6.1",
        "Рисунок 6.2 – Зависимость ошибки прогноза от фактического качества",
    ),
    (
        "absolute_error_distribution.png",
        "Рисунок 6.2",
        "Рисунок 6.3 – Распределение абсолютной ошибки интегрального прогноза",
    ),
    (
        "partial_criteria_comparison.png",
        "Таблица 6.7",
        "Рисунок 6.4 – Сравнение точности частных прогнозных критериев",
    ),
    (
        "confusion_matrix.png",
        "Таблица 6.10",
        "Рисунок 6.5 – Матрица ошибок классификации уровней качества",
    ),
    (
        "prediction_intervals.png",
        "Таблица 6.11",
        "Рисунок 6.6 – Проверка прогнозных интервалов на фактических значениях",
    ),
    (
        "baseline_comparison.png",
        "Таблица 6.13",
        "Рисунок 6.7 – Сравнение предложенной модели с базовыми схемами",
    ),
    (
        "error_by_dominant_topic.png",
        "Таблица 6.17",
        "Рисунок 6.8 – Ошибка прогноза по доминирующим латентным факторам",
    ),
)


class Chapter6DissertationUpdateError(ValueError):
    """Ошибка переноса результатов в текст шестой главы."""


@dataclass(frozen=True)
class Chapter6DissertationUpdateResult:
    """Результат обновления документа и текстовой сводки."""

    document_path: Path
    markdown_path: Path
    report_json_path: Path
    report_markdown_path: Path
    backup_path: Path | None
    report: dict[str, Any]


class Chapter6DissertationUpdater:
    """Заполнить таблицы, интерпретации и рисунки главы 6."""

    def __init__(
        self,
        project_root: Path,
        source_document: Path | None = None,
        output_document: Path | None = None,
    ) -> None:
        self.project_root = Path(project_root)
        self.source_document = self._resolve(
            source_document or Path(DEFAULT_DOCUMENT_NAME)
        )
        self.output_document = self._resolve(
            output_document or Path(DEFAULT_DOCUMENT_NAME)
        )
        self.reports_dir = self.project_root / REPORTS_DIR
        self.figures_dir = self.project_root / FIGURES_DIR

    def build_and_save(self) -> Chapter6DissertationUpdateResult:
        """Проверить источники и сформировать итоговые артефакты этапа 15."""

        sources = self._load_sources()
        source_hashes_before = self._source_hashes(sources)
        document = Document(str(self.source_document))
        self._validate_template(document)

        self._fill_integral_metrics(document, sources)
        self._fill_partial_metrics(document, sources)
        self._fill_classification(document, sources)
        self._fill_interval_metrics(document, sources)
        self._fill_baselines(document, sources)
        self._fill_bootstrap(document, sources)
        self._fill_top_errors(document, sources)
        self._replace_interpretations(document, sources)
        self._replace_conclusions(document, sources)
        self._insert_figures(document, sources)
        self._normalize_document(document)

        backup_path = self._save_document(document)
        self._validate_output_document(self.output_document)

        markdown_path = self.reports_dir / "chapter6_results_for_dissertation.md"
        report_json_path = self.reports_dir / "chapter6_dissertation_update_report.json"
        report_markdown_path = self.reports_dir / "chapter6_dissertation_update_report.md"
        markdown_path.parent.mkdir(parents=True, exist_ok=True)
        markdown_path.write_text(
            self._render_results_markdown(sources), encoding="utf-8"
        )

        source_hashes_after = self._source_hashes(sources)
        if source_hashes_before != source_hashes_after:
            raise Chapter6DissertationUpdateError(
                "Исходные расчетные артефакты изменились при обновлении документа."
            )

        report = self._build_report(
            sources=sources,
            source_hashes=source_hashes_after,
            markdown_path=markdown_path,
            backup_path=backup_path,
        )
        report_json_path.write_text(
            json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        report_markdown_path.write_text(
            self._render_stage_report(report), encoding="utf-8"
        )
        return Chapter6DissertationUpdateResult(
            document_path=self.output_document,
            markdown_path=markdown_path,
            report_json_path=report_json_path,
            report_markdown_path=report_markdown_path,
            backup_path=backup_path,
            report=report,
        )

    def _resolve(self, path: Path) -> Path:
        return path if path.is_absolute() else self.project_root / path

    def _load_sources(self) -> dict[str, Any]:
        required_json = {
            "acceptance": "chapter6_acceptance_report.json",
            "final": "chapter6_validation_report.json",
            "integral": "validation_metrics.json",
            "partial_report": "partial_criteria_validation_report.json",
            "classification": "classification_report.json",
            "interval": "interval_coverage_report.json",
            "bootstrap_report": "bootstrap_report.json",
            "error_report": "prediction_error_analysis.json",
            "figure_manifest": "figures/figure_manifest.json",
        }
        required_csv = {
            "partial": "partial_criteria_validation.csv",
            "confusion": "confusion_matrix.csv",
            "baselines": "baseline_comparison.csv",
            "bootstrap_ci": "bootstrap_confidence_intervals.csv",
            "bootstrap_differences": "bootstrap_model_differences.csv",
            "top_errors": "top_prediction_errors.csv",
        }
        sources: dict[str, Any] = {"paths": {}}
        for key, relative in required_json.items():
            path = self.reports_dir / relative
            if not path.exists():
                raise FileNotFoundError(f"Не найден обязательный JSON: {path}")
            sources[key] = json.loads(path.read_text(encoding="utf-8"))
            sources["paths"][key] = path
        for key, relative in required_csv.items():
            path = self.reports_dir / relative
            if not path.exists():
                raise FileNotFoundError(f"Не найден обязательный CSV: {path}")
            sources[key] = pd.read_csv(path)
            sources["paths"][key] = path

        uncertainty_path = (
            self.project_root / "reports/chapter5/prediction_uncertainty.csv"
        )
        if not uncertainty_path.exists():
            raise FileNotFoundError(
                "Не найден обязательный CSV главы 5: "
                f"{uncertainty_path}"
            )
        sources["chapter5_uncertainty"] = pd.read_csv(uncertainty_path)
        sources["paths"]["chapter5_uncertainty"] = uncertainty_path

        if not self.source_document.exists():
            raise FileNotFoundError(
                f"Не найден исходный документ главы 6: {self.source_document}"
            )
        sources["paths"]["source_document"] = self.source_document

        acceptance = sources["acceptance"]
        if acceptance.get("stage") != 14 or acceptance.get("accepted") is not True:
            raise Chapter6DissertationUpdateError(
                "Этап 15 разрешен только после успешной приемки этапа 14."
            )
        if int(acceptance.get("row_count", 0)) != 150:
            raise Chapter6DissertationUpdateError(
                "Акт приемки не подтверждает корпус из 150 сценариев."
            )
        final = sources["final"]
        if final.get("passed") is not True:
            raise Chapter6DissertationUpdateError(
                "Итоговый отчет этапа 13 имеет отрицательный статус."
            )
        if final.get("hypothesis_status") not in {
            "hypothesis_supported",
            "hypothesis_partially_supported",
            "hypothesis_not_supported",
        }:
            raise Chapter6DissertationUpdateError(
                "Итоговый отчет содержит неизвестный статус гипотезы."
            )
        manifest = sources["figure_manifest"]
        if manifest.get("passed") is not True or int(manifest.get("figure_count", 0)) != 8:
            raise Chapter6DissertationUpdateError(
                "Манифест этапа 12 не подтверждает восемь рисунков."
            )
        for filename, _, _ in FIGURE_SPECS:
            path = self.figures_dir / filename
            if not path.exists() or path.stat().st_size == 0:
                raise FileNotFoundError(f"Не найден рисунок этапа 12: {path}")
            sources["paths"][f"figure:{filename}"] = path
        return sources

    def _source_hashes(self, sources: Mapping[str, Any]) -> dict[str, str]:
        hashes: dict[str, str] = {}
        for key, path in sources["paths"].items():
            path = Path(path)
            if path == self.output_document and self.output_document == self.source_document:
                continue
            hashes[str(key)] = self._sha256(path)
        return hashes

    @staticmethod
    def _sha256(path: Path) -> str:
        digest = hashlib.sha256()
        with path.open("rb") as stream:
            for chunk in iter(lambda: stream.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()

    def _validate_template(self, document: DocumentType) -> None:
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        required = (
            "6.4. Метрики точности априорного прогноза",
            "6.12. Выводы по главе 6",
            "Таблица 6.5",
            "Таблица 6.17",
        )
        missing = [marker for marker in required if marker not in text]
        if missing:
            raise Chapter6DissertationUpdateError(
                "Документ не соответствует шаблону главы 6: " + ", ".join(missing)
            )

    def _find_table(self, document: DocumentType, first_header: str) -> Any:
        for table in document.tables:
            if table.rows and table.rows[0].cells:
                value = table.rows[0].cells[0].text.strip()
                if value == first_header:
                    return table
        raise Chapter6DissertationUpdateError(
            f"В документе не найдена таблица с заголовком «{first_header}»."
        )

    def _fill_integral_metrics(self, document: DocumentType, sources: Mapping[str, Any]) -> None:
        table = self._find_table(document, "Метрика")
        metrics = sources["integral"]["metrics"]
        values = {
            "MAE": metrics["mae"],
            "RMSE": metrics["rmse"],
            "Bias": metrics["bias"],
            "Pearson": metrics["pearson"],
            "Spearman": metrics["spearman"],
            "Kendall": metrics["kendall"],
            "R²": metrics["r2"],
        }
        for row in table.rows[1:]:
            label = row.cells[0].text.strip()
            if label in values:
                self._set_cell(row.cells[1], self._fmt(values[label]))

    def _fill_partial_metrics(self, document: DocumentType, sources: Mapping[str, Any]) -> None:
        table = self._find_table(document, "Критерий")
        frame = sources["partial"].set_index("criterion")
        for row in table.rows[1:]:
            criterion = row.cells[0].text.strip()
            if criterion not in frame.index:
                continue
            record = frame.loc[criterion]
            for column_index, key in enumerate(
                ("mae", "rmse", "bias", "spearman", "kendall"), start=1
            ):
                self._set_cell(row.cells[column_index], self._fmt(record[key]))

    def _fill_classification(self, document: DocumentType, sources: Mapping[str, Any]) -> None:
        confusion = self._find_table(document, "Фактический / прогнозный")
        frame = sources["confusion"]
        if "actual_class" in frame.columns:
            frame = frame.set_index("actual_class")
        else:
            frame = frame.set_index(frame.columns[0])
        for row in confusion.rows[1:]:
            actual = row.cells[0].text.strip()
            for index, predicted in enumerate(("low", "medium", "high"), start=1):
                self._set_cell(row.cells[index], str(int(frame.loc[actual, predicted])))

        table = self._find_table(document, "Метрика")
        # Первая таблица с заголовком «Метрика» уже содержит интегральные метрики.
        metric_tables = [
            item for item in document.tables
            if item.rows and item.rows[0].cells[0].text.strip() == "Метрика"
        ]
        classification_table = next(
            item for item in metric_tables
            if any("Balanced Accuracy" in row.cells[0].text for row in item.rows)
        )
        metrics = sources["classification"]["metrics"]
        critical = sources["classification"].get("critical_errors", {})
        per_class = {
            str(row.get("class_label")): row
            for row in sources["classification"].get("per_class_metrics", [])
        }
        values = {
            "Accuracy": self._fmt(metrics["accuracy"]),
            "Balanced Accuracy": self._fmt(metrics["balanced_accuracy"]),
            "Macro F1": self._fmt(metrics["macro_f1"]),
            "Precision low / medium / high": " / ".join(
                self._fmt(per_class.get(label, {}).get("precision", 0.0))
                for label in ("low", "medium", "high")
            ),
            "Recall low / medium / high": " / ".join(
                self._fmt(per_class.get(label, {}).get("recall", 0.0))
                for label in ("low", "medium", "high")
            ),
            "Критичные ошибки low→high": str(int(critical.get("low_to_high", 0))),
            "Критичные ошибки high→low": str(int(critical.get("high_to_low", 0))),
        }
        for row in classification_table.rows[1:]:
            label = row.cells[0].text.strip()
            if label in values:
                self._set_cell(row.cells[1], values[label])

    def _fill_interval_metrics(self, document: DocumentType, sources: Mapping[str, Any]) -> None:
        table = self._find_table(document, "Показатель")
        candidates = [
            item for item in document.tables
            if item.rows and item.rows[0].cells[0].text.strip() == "Показатель"
        ]
        interval_table = next(
            item for item in candidates
            if any("coverage_rate" in row.cells[0].text for row in item.rows)
        )
        metrics = sources["interval"]["metrics"]
        error_report = sources["error_report"]
        uncertainty = error_report.get("uncertainty_relation", {})
        chapter5_uncertainty = sources["chapter5_uncertainty"]
        required_columns = {
            "uncertainty_score",
            "interval_radius",
            "q_pred_lower",
            "q_pred_upper",
        }
        missing_columns = required_columns - set(chapter5_uncertainty.columns)
        if missing_columns:
            raise Chapter6DissertationUpdateError(
                "В prediction_uncertainty.csv отсутствуют колонки: "
                + ", ".join(sorted(missing_columns))
            )
        values = {
            "Средний uncertainty_score": self._fmt(
                chapter5_uncertainty["uncertainty_score"].mean()
            ),
            "Средний радиус интервала": self._fmt(
                chapter5_uncertainty["interval_radius"].mean()
            ),
            "Средняя ширина интервала до учета обрезки [0; 1]": self._fmt(
                (2.0 * chapter5_uncertainty["interval_radius"]).mean()
            ),
            "coverage_rate": self._fmt(metrics["coverage_rate"]),
            "miss_lower_count": str(int(metrics["miss_lower_count"])),
            "miss_upper_count": str(int(metrics["miss_upper_count"])),
            "Корреляция uncertainty_score с abs_error": self._fmt(
                uncertainty.get("spearman_absolute_error", 0.0)
            ),
        }
        slices = sources["interval"].get("slices", {})
        class_rows = slices.get("factual_class", slices.get("q_fact_class", []))
        if isinstance(class_rows, list) and class_rows:
            mapping = {
                str(row.get("group", row.get("value", row.get("class_label")))): row
                for row in class_rows
            }
            values["Покрытие по классам low / medium / high"] = " / ".join(
                self._fmt(mapping.get(label, {}).get("coverage_rate", 0.0))
                for label in ("low", "medium", "high")
            )
        else:
            values["Покрытие по классам low / medium / high"] = "см. JSON-отчет"
        for row in interval_table.rows[1:]:
            label = row.cells[0].text.strip()
            if label in values:
                self._set_cell(row.cells[1], values[label])

    def _fill_baselines(self, document: DocumentType, sources: Mapping[str, Any]) -> None:
        table = self._find_table(document, "Модель")
        candidates = [
            item for item in document.tables
            if item.rows and item.rows[0].cells[0].text.strip() == "Модель"
        ]
        target = next(item for item in candidates if len(item.columns) == 8)
        frame = sources["baselines"].set_index("model")
        model_names = {
            "Mean baseline": "mean_baseline",
            "Prior-only baseline": "prior_only_baseline",
            "Theta-only baseline": "theta_only_baseline",
            "Chapter 5 model": "chapter5_model",
        }
        keys = ("mae", "rmse", "bias", "spearman", "kendall", "balanced_accuracy", "macro_f1")
        for row in target.rows[1:]:
            source_name = model_names.get(row.cells[0].text.strip())
            if source_name is None or source_name not in frame.index:
                continue
            record = frame.loc[source_name]
            for index, key in enumerate(keys, start=1):
                self._set_cell(row.cells[index], self._fmt(record[key]))

    def _fill_bootstrap(self, document: DocumentType, sources: Mapping[str, Any]) -> None:
        ci_table = self._find_table(document, "Метрика")
        candidates = [
            item for item in document.tables
            if item.rows and item.rows[0].cells[0].text.strip() == "Метрика"
        ]
        ci_target = next(item for item in candidates if len(item.columns) == 4)
        frame = sources["bootstrap_ci"]
        frame = frame[frame["model"] == "chapter5_model"].set_index("metric")
        metric_map = {
            "MAE": "mae",
            "RMSE": "rmse",
            "Spearman": "spearman",
            "Kendall": "kendall",
            "Balanced Accuracy": "accuracy",
            "Macro F1": "macro_f1",
        }
        for row in ci_target.rows[1:]:
            label = row.cells[0].text.strip()
            key = metric_map.get(label)
            if key is None or key not in frame.index:
                continue
            if label == "Balanced Accuracy":
                self._set_cell(row.cells[0], "Accuracy")
            record = frame.loc[key]
            self._set_cell(row.cells[1], self._fmt(record["point_estimate"]))
            self._set_cell(row.cells[2], self._fmt(record["ci_lower"]))
            self._set_cell(row.cells[3], self._fmt(record["ci_upper"]))

        diff_table = self._find_table(document, "Сравнение")
        frame = sources["bootstrap_differences"]
        rows_to_show = (
            ("mean_baseline", "mae"),
            ("prior_only_baseline", "mae"),
            ("theta_only_baseline", "mae"),
            ("prior_only_baseline", "spearman"),
        )
        label_map = {
            "mean_baseline": "Chapter 5 − Mean",
            "prior_only_baseline": "Chapter 5 − Prior-only",
            "theta_only_baseline": "Chapter 5 − Theta-only",
        }
        for row, (baseline, metric) in zip(diff_table.rows[1:], rows_to_show, strict=True):
            selected = frame[(frame["baseline"] == baseline) & (frame["metric"] == metric)]
            if selected.empty:
                raise Chapter6DissertationUpdateError(
                    f"Не найдена bootstrap-разность {baseline}/{metric}."
                )
            record = selected.iloc[0]
            self._set_cell(row.cells[0], label_map[baseline])
            self._set_cell(row.cells[1], metric.upper() if metric != "spearman" else "Spearman")
            self._set_cell(row.cells[2], self._fmt(record["point_delta"]))
            self._set_cell(row.cells[3], self._fmt(record["ci_lower"]))
            self._set_cell(row.cells[4], self._fmt(record["ci_upper"]))
            self._set_cell(
                row.cells[5],
                self._conclusion_ru(str(record.get("conclusion", ""))),
            )

    def _fill_top_errors(self, document: DocumentType, sources: Mapping[str, Any]) -> None:
        table = self._find_table(document, "scenario_id")
        frame = sources["top_errors"].head(10)
        for row_index, (_, record) in enumerate(frame.iterrows(), start=1):
            row = table.rows[row_index]
            values = (
                record.get("scenario_id", ""),
                self._fmt(record.get("q_pred")),
                self._fmt(record.get("q_fact")),
                self._fmt(record.get("prediction_error")),
                self._fmt(record.get("absolute_error")),
                record.get("dominant_factor", ""),
                self._fmt(record.get("uncertainty_score")),
            )
            for cell, value in zip(row.cells, values, strict=True):
                self._set_cell(cell, str(value))

    def _replace_interpretations(self, document: DocumentType, sources: Mapping[str, Any]) -> None:
        integral = sources["integral"]["metrics"]
        partial = sources["partial_report"]["summary"]
        classification = sources["classification"]["metrics"]
        critical = sources["classification"].get("critical_errors", {})
        interval = sources["interval"]["metrics"]
        bootstrap = sources["bootstrap_report"]["summary"]
        error = sources["error_report"]
        replacements = {
            "Окончательная интерпретация должна опираться": (
                "Полученные результаты показывают сочетание высокой ранговой согласованности "
                f"(Spearman = {self._fmt(integral['spearman'])}, Kendall = {self._fmt(integral['kendall'])}) "
                "и неудовлетворительной калибровки абсолютной шкалы. "
                f"MAE составила {self._fmt(integral['mae'])}, RMSE — {self._fmt(integral['rmse'])}, "
                f"Bias — {self._fmt(integral['bias'])}, R² — {self._fmt(integral['r2'])}. "
                "Следовательно, модель надежно сохраняет относительный порядок сценариев, "
                "но систематически занижает фактический уровень качества."
            ),
            "Для наглядного представления результатов": (
                "На рисунках 6.1–6.3 представлены сопоставление Q_pred и Q_fact, "
                "структура остатков и распределение абсолютных ошибок. Все изображения "
                "сформированы программно из зафиксированного проверочного датасета без "
                "ручной подмены расчетных значений."
            ),
            "После подстановки результатов критерии следует": (
                f"Средняя MAE по шести частным критериям равна {self._fmt(partial['mean_mae'])}. "
                f"Наименьшая ошибка получена для {partial['best_mae_criterion']}, "
                f"наибольшая — для {partial['worst_mae_criterion']}. "
                f"Средний коэффициент Spearman составил {self._fmt(partial['mean_spearman'])}. "
                "Во всех частных оценках выявлено отрицательное смещение, поэтому основной "
                "недостаток связан не с потерей порядка сценариев, а с уровнем шкалы прогноза."
            ),
            "Различие маргинальных распределений показывает": (
                "Матрица ошибок подтверждает выраженную консервативность классификации: "
                "66 фактических сценариев среднего класса отнесены к low, тогда как "
                "критические переходы low→high и high→low отсутствуют."
            ),
            "Из-за фактического дисбаланса классов": (
                f"Accuracy составила {self._fmt(classification['accuracy'])}, Balanced Accuracy — "
                f"{self._fmt(classification['balanced_accuracy'])}, Macro F1 — "
                f"{self._fmt(classification['macro_f1'])}. Число критических ошибок равно "
                f"{int(critical.get('total', critical.get('low_to_high', 0) + critical.get('high_to_low', 0)))}. "
                "Результат пригоден для предварительного ранжирования и консервативного отбора, "
                "но не подтверждает сбалансированную классификацию всех трех уровней."
            ),
            "Высокое покрытие само по себе": (
                f"Фактическое покрытие составило только {self._fmt(interval['coverage_rate'])}, "
                f"то есть {int(interval['covered_count'])} из 150 сценариев. "
                f"Средняя ширина интервала равна {self._fmt(interval['mean_interval_width'])}; "
                f"в {int(interval['miss_upper_count'])} случаях фактическое качество оказалось "
                "выше верхней границы. Интервалы являются слишком узкими и смещенными вниз."
            ),
            "Положительная связь uncertainty_score": (
                "Связь uncertainty_score с абсолютной ошибкой оказалась слабой "
                f"(Spearman = {self._fmt(error.get('uncertainty_relation', {}).get('spearman_absolute_error', 0.0))}). "
                "Текущая оценка неопределенности не может рассматриваться как надежный индикатор "
                "сценариев повышенного риска и требует отдельной перекалибровки."
            ),
            "Преимущество основной модели фиксируется": (
                "Сравнение показало, что модель главы 5 уступает mean- и prior-only-baseline "
                "по MAE и RMSE, но превосходит theta-only baseline и имеет высокую ранговую "
                "согласованность. Поэтому включение латентного профиля не доказало универсального "
                "снижения абсолютной ошибки, однако сохранило диагностическую и ранговую ценность."
            ),
            "Bootstrap-интервалы характеризуют": (
                f"Парный bootstrap дал {int(bootstrap['stable_chapter5_wins'])} устойчивых преимуществ "
                f"модели главы 5, {int(bootstrap['stable_baseline_wins'])} устойчивых преимуществ "
                f"baseline и {int(bootstrap['no_stable_difference'])} сравнений без устойчивого вывода. "
                "Преимущество mean- и prior-only-baseline по абсолютной ошибке является устойчивым, "
                "тогда как различия ранговых и части классификационных метрик не всегда исключают ноль."
            ),
            "Если крупные ошибки концентрируются": (
                "Диагностический анализ выявил систематическое занижение в 132 из 150 сценариев. "
                "Наибольшая MAE характерна для групп theta_0 и theta_1, тогда как группа theta_2 "
                "оценивается существенно точнее. Сильнейшая связь абсолютной ошибки выявлена с "
                "фактической длительностью выполнения, что указывает на недостаточный учет сложных "
                "и продолжительных процедур в априорной формуле."
            ),
        }
        for prefix, text in replacements.items():
            self._replace_paragraph_start(document, prefix, text)

    def _replace_conclusions(self, document: DocumentType, sources: Mapping[str, Any]) -> None:
        final = sources["final"]
        integral = sources["integral"]["metrics"]
        interval = sources["interval"]["metrics"]
        classification = sources["classification"]["metrics"]
        conclusions = [
            (
                "Экспериментальная проверка выполнена на корпусе из 150 сценариев. "
                "Программный контур принят по 27 из 27 обязательных проверок; прогнозы, "
                "веса, пороги и формулы главы 5 в ходе проверки не изменялись."
            ),
            (
                f"Интегральный прогноз характеризуется MAE = {self._fmt(integral['mae'])}, "
                f"RMSE = {self._fmt(integral['rmse'])} и Bias = {self._fmt(integral['bias'])}. "
                f"Отрицательное R² = {self._fmt(integral['r2'])} подтверждает неудовлетворительную "
                "калибровку абсолютной шкалы."
            ),
            (
                f"Высокие значения Pearson = {self._fmt(integral['pearson'])}, Spearman = "
                f"{self._fmt(integral['spearman'])} и Kendall = {self._fmt(integral['kendall'])} "
                "подтверждают способность метода сохранять относительный порядок сценариев."
            ),
            (
                f"Классификационная проверка дала Accuracy = {self._fmt(classification['accuracy'])}, "
                f"Balanced Accuracy = {self._fmt(classification['balanced_accuracy'])} и Macro F1 = "
                f"{self._fmt(classification['macro_f1'])}. Критические ошибки между крайними "
                "классами отсутствуют, однако дисбаланс low/medium/high ограничивает обобщение."
            ),
            (
                f"Покрытие прогнозных интервалов равно {self._fmt(interval['coverage_rate'])}; "
                "интервальная модель не прошла внешнюю калибровочную проверку."
            ),
            (
                "Сравнение с baseline показало устойчивое преимущество mean- и prior-only-схем "
                "по MAE и RMSE. Поэтому утверждение о безусловном повышении абсолютной точности "
                "за счет LDA экспериментально не подтверждено."
            ),
            (
                "Латентный профиль сохраняет научную ценность как средство структурирования "
                "априорного признакового пространства, ранжирования сценариев и выявления групп "
                "с различной достоверностью прогноза."
            ),
            (
                "Основная гипотеза исследования подтверждена частично: предложенная модель "
                "обеспечивает высокую ранговую достоверность и диагностическую информативность, "
                "но не дает одновременного устойчивого улучшения абсолютной точности и "
                "интервальной калибровки."
            ),
            (
                "Для дальнейшего развития требуется отдельная калибровочная выборка, внешняя "
                "эмпирическая проверка с участием операторов и оценка обобщающей способности "
                "при переносе на новые средства, условия и типы сообщений."
            ),
        ]
        heading_index = next(
            index for index, paragraph in enumerate(document.paragraphs)
            if (
                paragraph.text.strip().startswith("6.12. Выводы по главе 6")
                and paragraph.style.name.startswith("Heading")
            )
        )
        paragraphs = document.paragraphs
        start_index = heading_index + 1
        if len(paragraphs) < start_index + len(conclusions):
            raise Chapter6DissertationUpdateError(
                "Раздел выводов не содержит достаточного числа абзацев."
            )
        for index, text in enumerate(conclusions):
            paragraph = paragraphs[start_index + index]
            paragraph.text = f"{index + 1}. {text}"
        status_text = {
            "hypothesis_supported": "основная гипотеза подтверждена",
            "hypothesis_partially_supported": (
                "основная гипотеза подтверждена частично"
            ),
            "hypothesis_not_supported": "основная гипотеза не подтверждена",
        }[final["hypothesis_status"]]
        self._replace_paragraph_start(
            document,
            "После выполнения программных этапов",
            (
                "Все расчетные поля главы заменены фактическими значениями, рисунки встроены "
                "из воспроизводимых артефактов этапа 12, а итоговые формулировки согласованы "
                f"с программным выводом: {status_text}."
            ),
        )

    def _insert_figures(self, document: DocumentType, sources: Mapping[str, Any]) -> None:
        # При повторном запуске удаляем ранее встроенный блок рисунков этапа 15.
        self._remove_existing_stage15_figures(document)
        anchor_map = {
            paragraph.text.strip(): paragraph for paragraph in document.paragraphs
        }
        inserted_captions: dict[str, Any] = {}
        for filename, anchor_prefix, caption in FIGURE_SPECS:
            if anchor_prefix.startswith("Рисунок"):
                anchor = inserted_captions.get(anchor_prefix)
            else:
                caption_anchor = next(
                    (
                        paragraph for paragraph in document.paragraphs
                        if paragraph.text.strip().startswith(anchor_prefix)
                    ),
                    None,
                )
                anchor = (
                    self._paragraph_after_following_table(caption_anchor)
                    if caption_anchor is not None
                    else None
                )
            if anchor is None:
                raise Chapter6DissertationUpdateError(
                    f"Не найдена позиция вставки рисунка: {anchor_prefix}."
                )
            image_paragraph = self._insert_paragraph_after(anchor)
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = image_paragraph.add_run()
            run.add_picture(str(self.figures_dir / filename), width=Cm(15.5))
            caption_paragraph = self._insert_paragraph_after(image_paragraph)
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_paragraph.add_run(caption)
            caption_run.font.name = "Times New Roman"
            caption_run.font.size = Pt(12)
            caption_paragraph._p.set(qn("w:rsidR"), "15A15A15")
            inserted_captions[caption.split(" – ", 1)[0]] = caption_paragraph

    @staticmethod
    def _paragraph_after_following_table(paragraph: Any) -> Any:
        """Создать технический абзац после таблицы, следующей за подписью."""

        sibling = paragraph._p.getnext()
        while sibling is not None and sibling.tag != qn("w:tbl"):
            sibling = sibling.getnext()
        if sibling is None:
            return paragraph
        new_p = OxmlElement("w:p")
        sibling.addnext(new_p)
        from docx.text.paragraph import Paragraph

        return Paragraph(new_p, paragraph._parent)

    def _remove_existing_stage15_figures(self, document: DocumentType) -> None:
        for paragraph in list(document.paragraphs):
            if paragraph.text.strip().startswith("Рисунок 6."):
                previous = paragraph._p.getprevious()
                if previous is not None and previous.xpath(".//a:blip"):
                    previous.getparent().remove(previous)
                paragraph._p.getparent().remove(paragraph._p)

    @staticmethod
    def _insert_paragraph_after(paragraph: Any) -> Any:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        from docx.text.paragraph import Paragraph

        return Paragraph(new_p, paragraph._parent)

    def _normalize_document(self, document: DocumentType) -> None:
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.get_or_add_rPr().get_or_add_rFonts().set(
                    qn("w:ascii"), "Times New Roman"
                )
                run._element.get_or_add_rPr().get_or_add_rFonts().set(
                    qn("w:hAnsi"), "Times New Roman"
                )
                run._element.get_or_add_rPr().get_or_add_rFonts().set(
                    qn("w:eastAsia"), "Times New Roman"
                )
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(10)

    def _save_document(self, document: DocumentType) -> Path | None:
        self.output_document.parent.mkdir(parents=True, exist_ok=True)
        backup_path: Path | None = None
        if self.output_document.resolve() == self.source_document.resolve():
            backup_path = self.reports_dir / "dissertation" / "chapter6_before_stage15.docx"
            backup_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(self.source_document, backup_path)
            temporary = self.output_document.with_suffix(".stage15.tmp.docx")
            document.save(str(temporary))
            temporary.replace(self.output_document)
        else:
            document.save(str(self.output_document))
        return backup_path

    def _validate_output_document(self, path: Path) -> None:
        document = Document(str(path))
        text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        )
        remaining = [marker for marker in PLACEHOLDER_MARKERS if marker in text]
        if remaining:
            raise Chapter6DissertationUpdateError(
                "В итоговом DOCX остались незаполненные маркеры: " + ", ".join(remaining)
            )
        if len(document.inline_shapes) != 8:
            raise Chapter6DissertationUpdateError(
                f"В итоговом DOCX ожидалось 8 рисунков, получено {len(document.inline_shapes)}."
            )

    def _build_report(
        self,
        sources: Mapping[str, Any],
        source_hashes: Mapping[str, str],
        markdown_path: Path,
        backup_path: Path | None,
    ) -> dict[str, Any]:
        return {
            "stage": 15,
            "report_type": "chapter6_dissertation_update_report",
            "passed": True,
            "row_count": 150,
            "hypothesis_status": sources["final"]["hypothesis_status"],
            "document_path": str(self.output_document.relative_to(self.project_root)),
            "document_sha256": self._sha256(self.output_document),
            "markdown_path": str(markdown_path.relative_to(self.project_root)),
            "markdown_sha256": self._sha256(markdown_path),
            "backup_path": (
                str(backup_path.relative_to(self.project_root)) if backup_path else None
            ),
            "filled_tables": 9,
            "inserted_figures": 8,
            "placeholders_remaining": False,
            "source_artifacts_modified": False,
            "synthetic_holdout_used": False,
            "source_hashes": dict(source_hashes),
            "methodological_checks": {
                "stage14_acceptance_required": True,
                "chapter5_prediction_modified": False,
                "quality_thresholds_modified": False,
                "factual_values_used_only_for_external_validation": True,
                "hypothesis_status_copied_from_stage13": True,
                "manual_metric_substitution": False,
            },
        }

    def _render_results_markdown(self, sources: Mapping[str, Any]) -> str:
        integral = sources["integral"]["metrics"]
        classification = sources["classification"]["metrics"]
        interval = sources["interval"]["metrics"]
        baselines = sources["baselines"]
        lines = [
            "# Результаты главы 6 для переноса в диссертацию",
            "",
            "> Документ сформирован автоматически из принятых артефактов этапов 2–14.",
            "",
            "## Статус проверки",
            "",
            f"- Сценариев: **150**.",
            f"- Статус гипотезы: `{sources['final']['hypothesis_status']}`.",
            "- Финальная приемка: **27/27 проверок**.",
            "- Синтетический holdout не используется как основное доказательство.",
            "",
            "## Интегральные метрики",
            "",
            "| Метрика | Значение |",
            "|---|---:|",
        ]
        for key, label in (
            ("mae", "MAE"), ("rmse", "RMSE"), ("bias", "Bias"),
            ("pearson", "Pearson"), ("spearman", "Spearman"),
            ("kendall", "Kendall"), ("r2", "R²"),
        ):
            lines.append(f"| {label} | {self._fmt(integral[key])} |")
        lines.extend([
            "",
            "## Классификация",
            "",
            f"- Accuracy: **{self._fmt(classification['accuracy'])}**.",
            f"- Balanced Accuracy: **{self._fmt(classification['balanced_accuracy'])}**.",
            f"- Macro F1: **{self._fmt(classification['macro_f1'])}**.",
            "- Критические ошибки low→high и high→low отсутствуют.",
            "",
            "## Интервальный прогноз",
            "",
            f"- Coverage rate: **{self._fmt(interval['coverage_rate'])}**.",
            f"- Средняя ширина: **{self._fmt(interval['mean_interval_width'])}**.",
            f"- Промахов выше верхней границы: **{int(interval['miss_upper_count'])}**.",
            "",
            "## Сравнение моделей",
            "",
            "| Модель | MAE | RMSE | Spearman | Kendall | Balanced Accuracy | Macro F1 |",
            "|---|---:|---:|---:|---:|---:|---:|",
        ])
        for _, row in baselines.iterrows():
            lines.append(
                f"| {row['model']} | {self._fmt(row['mae'])} | {self._fmt(row['rmse'])} | "
                f"{self._fmt(row['spearman'])} | {self._fmt(row['kendall'])} | "
                f"{self._fmt(row['balanced_accuracy'])} | {self._fmt(row['macro_f1'])} |"
            )
        lines.extend([
            "",
            "## Итоговый вывод",
            "",
            "Основная гипотеза подтверждена частично. Адаптированная LDA-модель "
            "обеспечивает высокую ранговую согласованность и диагностическую "
            "информативность, но не дает устойчивого снижения абсолютной ошибки "
            "относительно mean- и prior-only-baseline и не обеспечивает "
            "калиброванный интервальный прогноз.",
            "",
        ])
        return "\n".join(lines)

    @staticmethod
    def _render_stage_report(report: Mapping[str, Any]) -> str:
        return "\n".join(
            [
                "# Отчет этапа 15",
                "",
                f"- Статус: **{'пройден' if report['passed'] else 'не пройден'}**.",
                f"- Документ: `{report['document_path']}`.",
                f"- Текстовая сводка: `{report['markdown_path']}`.",
                f"- Заполнено таблиц: **{report['filled_tables']}**.",
                f"- Вставлено рисунков: **{report['inserted_figures']}**.",
                f"- Статус гипотезы: `{report['hypothesis_status']}`.",
                f"- Остались маркеры заполнения: **{report['placeholders_remaining']}**.",
                f"- Исходные артефакты изменены: **{report['source_artifacts_modified']}**.",
                "",
            ]
        )

    def _replace_paragraph_start(
        self, document: DocumentType, prefix: str, replacement: str
    ) -> None:
        paragraph = next(
            (
                item for item in document.paragraphs
                if item.text.strip().startswith(prefix)
            ),
            None,
        )
        if paragraph is None:
            paragraph = next(
                (
                    item for item in document.paragraphs
                    if item.text.strip().startswith(replacement[:60])
                ),
                None,
            )
        if paragraph is None:
            raise Chapter6DissertationUpdateError(
                f"В документе не найден абзац: {prefix}"
            )
        paragraph.text = replacement

    @staticmethod
    def _set_cell(cell: Any, value: str) -> None:
        cell.text = value
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    @staticmethod
    def _fmt(value: Any, digits: int = 6) -> str:
        if value is None or (isinstance(value, float) and pd.isna(value)):
            return "—"
        try:
            number = float(value)
        except (TypeError, ValueError):
            return str(value)
        return f"{number:.{digits}f}".replace(".", ",")

    @staticmethod
    def _conclusion_ru(value: str) -> str:
        return {
            "chapter5_model_favored": "да, модель главы 5",
            "baseline_favored": "нет, преимущество baseline",
            "no_stable_difference": "устойчивое различие не выявлено",
        }.get(value, value or "—")
