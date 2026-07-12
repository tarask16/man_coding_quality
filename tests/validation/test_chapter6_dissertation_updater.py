"""Тесты этапа 15: перенос результатов в текст диссертации."""

from __future__ import annotations

import base64
import json
from pathlib import Path

import pandas as pd
import pytest
from docx import Document

from manual_coding_sim.validation.chapter6_dissertation_updater import (
    Chapter6DissertationUpdateError,
    Chapter6DissertationUpdater,
    DEFAULT_DOCUMENT_NAME,
)
from manual_coding_sim.validation.chapter6_runner import main

PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9Z2qQAAAAASUVORK5CYII="
)


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value


def _write_template(path: Path) -> None:
    document = Document()
    document.add_paragraph("ГЛАВА 6")
    document.add_heading("6.4. Метрики точности априорного прогноза", level=1)
    document.add_paragraph("Окончательная интерпретация должна опираться на сравнение.")
    document.add_paragraph("Таблица 6.5 - Результаты проверки интегрального прогноза")
    _add_table(
        document,
        ["Метрика", "Значение", "Интерпретация"],
        [[name, "[рассчитать]", "описание"] for name in [
            "MAE", "RMSE", "Bias", "Pearson", "Spearman", "Kendall", "R²"
        ]],
    )
    document.add_paragraph("Для наглядного представления результатов рекомендуется включить графики.")
    document.add_heading("6.5. Проверка частных прогнозных критериев", level=1)
    document.add_paragraph("Таблица 6.7 - Шаблон результатов проверки частных критериев")
    _add_table(
        document,
        ["Критерий", "MAE", "RMSE", "Bias", "Spearman", "Kendall"],
        [[name, "[рассчитать]", "[рассчитать]", "[рассчитать]", "[рассчитать]", "[рассчитать]"]
         for name in ["q_acc", "q_time", "q_effort", "q_res", "q_rep", "q_fit"]],
    )
    document.add_paragraph("После подстановки результатов критерии следует ранжировать.")
    document.add_heading("6.6. Проверка классификации уровней качества", level=1)
    document.add_paragraph("Различие маргинальных распределений показывает предварительный сигнал.")
    document.add_paragraph("Таблица 6.9 - Матрица ошибок классификации качества")
    _add_table(
        document,
        ["Фактический / прогнозный", "low", "medium", "high"],
        [[label, "[n]", "[n]", "[n]"] for label in ["low", "medium", "high"]],
    )
    document.add_paragraph("Таблица 6.10 - Метрики классификационной проверки")
    _add_table(
        document,
        ["Метрика", "Значение", "Назначение"],
        [[name, "[рассчитать]", "описание"] for name in [
            "Accuracy", "Balanced Accuracy", "Macro F1",
            "Precision low / medium / high", "Recall low / medium / high",
            "Критичные ошибки low→high", "Критичные ошибки high→low",
        ]],
    )
    document.add_paragraph("Из-за фактического дисбаланса классов обычная Accuracy недостаточна.")
    document.add_heading("6.7. Проверка интервального прогноза", level=1)
    document.add_paragraph("Таблица 6.11 - Проверка интервального прогноза")
    _add_table(
        document,
        ["Показатель", "Значение"],
        [[name, "[рассчитать]"] for name in [
            "Средний uncertainty_score", "Средний радиус интервала",
            "Средняя ширина интервала до учета обрезки [0; 1]", "coverage_rate",
            "miss_lower_count", "miss_upper_count",
            "Покрытие по классам low / medium / high",
            "Корреляция uncertainty_score с abs_error",
        ]],
    )
    document.add_paragraph("Высокое покрытие само по себе не является достаточным результатом.")
    document.add_paragraph("Положительная связь uncertainty_score с абсолютной ошибкой подтверждает ценность.")
    document.add_heading("6.8. Сравнение с базовыми моделями без LDA", level=1)
    document.add_paragraph("Таблица 6.13 - Итоговое сравнение моделей")
    _add_table(
        document,
        ["Модель", "MAE", "RMSE", "Bias", "Spearman", "Kendall", "Balanced Acc.", "Macro F1"],
        [[name] + ["[рассчитать]"] * 7 for name in [
            "Mean baseline", "Prior-only baseline", "Theta-only baseline", "Chapter 5 model"
        ]],
    )
    document.add_paragraph("Преимущество основной модели фиксируется только при согласованной картине.")
    document.add_heading("6.9. Статистическая устойчивость результатов", level=1)
    document.add_paragraph("Таблица 6.14 - Bootstrap-доверительные интервалы основной модели")
    _add_table(
        document,
        ["Метрика", "Точечная оценка", "95% CI lower", "95% CI upper"],
        [[name, "[рассчитать]", "[рассчитать]", "[рассчитать]"] for name in [
            "MAE", "RMSE", "Spearman", "Kendall", "Balanced Accuracy", "Macro F1"
        ]],
    )
    document.add_paragraph("Таблица 6.15 - Парные bootstrap-разности модели главы 5 и baseline")
    _add_table(
        document,
        ["Сравнение", "Метрика", "Δ", "95% CI lower", "95% CI upper", "Устойчивое преимущество"],
        [["x", "x", "[рассчитать]", "[рассчитать]", "[рассчитать]", "[да/нет]"] for _ in range(4)],
    )
    document.add_paragraph("Bootstrap-интервалы характеризуют устойчивость результатов.")
    document.add_heading("6.10. Анализ ошибок прогноза", level=1)
    document.add_paragraph("Если крупные ошибки концентрируются в одном латентном факторе.")
    document.add_paragraph("Таблица 6.17 - Топ-10 сценариев по абсолютной ошибке")
    _add_table(
        document,
        ["scenario_id", "Q_pred", "Q_fact", "error", "abs_error", "topic", "uncertainty"],
        [["[автозаполнение]"] * 7 for _ in range(10)],
    )
    document.add_heading("6.12. Выводы по главе 6", level=1)
    for index in range(9):
        document.add_paragraph(
            "В настоящем черновике сформирована полная методическая структура."
            if index == 0 else f"{index + 1}. Черновой вывод."
        )
    document.add_paragraph("Таблица 6.18 - Артефакты для автоматической финализации главы")
    document.add_paragraph("После выполнения программных этапов поля подлежат подстановке.")
    document.save(path)


def _write_project(root: Path) -> None:
    reports = root / "reports/chapter6"
    figures = reports / "figures"
    figures.mkdir(parents=True)
    _write_template(root / DEFAULT_DOCUMENT_NAME)

    acceptance = {
        "stage": 14, "accepted": True, "row_count": 150,
        "check_count": 27, "passed_check_count": 27,
    }
    final = {
        "stage": 13, "passed": True, "row_count": 150,
        "hypothesis_status": "hypothesis_partially_supported",
    }
    integral = {
        "stage": 5, "passed": True, "row_count": 150,
        "metrics": {
            "mae": 0.159244, "rmse": 0.194403, "bias": -0.149399,
            "median_absolute_error": 0.14524, "max_absolute_error": 0.376663,
            "pearson": 0.892053, "spearman": 0.881769, "kendall": 0.699329,
            "r2": -1.971612, "q_pred_mean": 0.498844, "q_fact_mean": 0.648242,
        },
    }
    partial_report = {
        "stage": 6, "passed": True, "row_count": 150,
        "summary": {
            "mean_mae": 0.218271, "mean_spearman": 0.694796,
            "best_mae_criterion": "q_effort", "worst_mae_criterion": "q_fit",
        },
    }
    classification = {
        "stage": 7, "passed": True, "row_count": 150,
        "metrics": {
            "accuracy": 0.446667, "balanced_accuracy": 0.667823,
            "macro_f1": 0.422515, "weighted_f1": 0.555257,
        },
        "critical_errors": {"low_to_high": 0, "high_to_low": 0, "total": 0},
        "per_class_metrics": [
            {"class_label": "low", "precision": 0.014925, "recall": 1.0},
            {"class_label": "medium", "precision": 0.673469, "recall": 0.33},
            {"class_label": "high", "precision": 0.970588, "recall": 0.673469},
        ],
    }
    interval = {
        "stage": 8, "passed": True, "row_count": 150,
        "metrics": {
            "coverage_rate": 0.186667, "covered_count": 28,
            "miss_count": 122, "mean_interval_width": 0.084183,
            "median_interval_width": 0.097432, "miss_lower_count": 11,
            "miss_upper_count": 111, "mean_distance_to_interval": 0.121179,
        },
        "slices": {},
    }
    bootstrap_report = {
        "stage": 10, "passed": True, "row_count": 150,
        "summary": {
            "stable_chapter5_wins": 7, "stable_baseline_wins": 5,
            "no_stable_difference": 6, "comparison_count": 18,
        },
    }
    error_report = {
        "stage": 11, "passed": True, "row_count": 150,
        "uncertainty_relation": {"spearman_absolute_error": 0.164670},
    }
    json_payloads = {
        "chapter6_acceptance_report.json": acceptance,
        "chapter6_validation_report.json": final,
        "validation_metrics.json": integral,
        "partial_criteria_validation_report.json": partial_report,
        "classification_report.json": classification,
        "interval_coverage_report.json": interval,
        "bootstrap_report.json": bootstrap_report,
        "prediction_error_analysis.json": error_report,
    }
    for name, payload in json_payloads.items():
        (reports / name).write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

    partial = []
    for index, criterion in enumerate(["q_acc", "q_time", "q_effort", "q_res", "q_rep", "q_fit"]):
        partial.append({
            "criterion": criterion, "mae": 0.15 + index * 0.01,
            "rmse": 0.20 + index * 0.01, "bias": -0.10 - index * 0.01,
            "spearman": 0.60 + index * 0.03, "kendall": 0.45 + index * 0.02,
        })
    pd.DataFrame(partial).to_csv(reports / "partial_criteria_validation.csv", index=False)
    pd.DataFrame(
        [[1, 0, 0], [66, 33, 1], [0, 16, 33]],
        index=["low", "medium", "high"], columns=["low", "medium", "high"],
    ).rename_axis("actual_class").to_csv(reports / "confusion_matrix.csv")

    baselines = [
        ["mean_baseline", 0.097227, 0.112903, 0.0, -0.071744, -0.051253, 0.333333, 0.266667],
        ["prior_only_baseline", 0.111822, 0.132113, -0.107710, 0.859052, 0.666577, 0.645646, 0.393155],
        ["theta_only_baseline", 0.304125, 0.360544, -0.228721, 0.866689, 0.677494, 0.624898, 0.343074],
        ["chapter5_model", 0.159244, 0.194403, -0.149399, 0.881769, 0.699329, 0.667823, 0.422515],
    ]
    pd.DataFrame(baselines, columns=[
        "model", "mae", "rmse", "bias", "spearman", "kendall", "balanced_accuracy", "macro_f1"
    ]).to_csv(reports / "baseline_comparison.csv", index=False)

    ci_rows = []
    for metric, point in {
        "mae": 0.159244, "rmse": 0.194403, "spearman": 0.881769,
        "kendall": 0.699329, "accuracy": 0.446667, "macro_f1": 0.422515,
    }.items():
        ci_rows.append({
            "model": "chapter5_model", "metric": metric, "point_estimate": point,
            "ci_lower": point - 0.03, "ci_upper": point + 0.03,
        })
    pd.DataFrame(ci_rows).to_csv(reports / "bootstrap_confidence_intervals.csv", index=False)
    differences = []
    for baseline in ["mean_baseline", "prior_only_baseline", "theta_only_baseline"]:
        for metric in ["mae", "rmse", "spearman", "kendall", "accuracy", "macro_f1"]:
            conclusion = "no_stable_difference"
            if metric in {"mae", "rmse"} and baseline in {"mean_baseline", "prior_only_baseline"}:
                conclusion = "baseline_favored"
            elif baseline == "theta_only_baseline":
                conclusion = "chapter5_model_favored"
            differences.append({
                "baseline": baseline, "metric": metric, "point_delta": 0.05,
                "ci_lower": 0.01, "ci_upper": 0.09, "conclusion": conclusion,
            })
    pd.DataFrame(differences).to_csv(reports / "bootstrap_model_differences.csv", index=False)

    top = []
    for index in range(10):
        top.append({
            "scenario_id": f"S{index + 1:03d}", "q_pred": 0.3 + index * 0.01,
            "q_fact": 0.7 + index * 0.01, "prediction_error": -0.4,
            "absolute_error": 0.4, "dominant_factor": f"theta_{index % 3}",
            "uncertainty_score": 0.2 + index * 0.01,
        })
    pd.DataFrame(top).to_csv(reports / "top_prediction_errors.csv", index=False)

    chapter5_reports = root / "reports/chapter5"
    chapter5_reports.mkdir(parents=True, exist_ok=True)
    pd.DataFrame({
        "uncertainty_score": [0.20, 0.30],
        "interval_radius": [0.04, 0.05],
        "q_pred_lower": [0.40, 0.50],
        "q_pred_upper": [0.48, 0.60],
    }).to_csv(chapter5_reports / "prediction_uncertainty.csv", index=False)

    manifest = {
        "stage": 12, "passed": True, "row_count": 150,
        "figure_count": 8, "dpi": 300,
    }
    (figures / "figure_manifest.json").write_text(json.dumps(manifest), encoding="utf-8")
    for filename in [
        "q_pred_vs_q_fact.png", "residuals_vs_q_fact.png",
        "absolute_error_distribution.png", "confusion_matrix.png",
        "baseline_comparison.png", "prediction_intervals.png",
        "error_by_dominant_topic.png", "partial_criteria_comparison.png",
    ]:
        (figures / filename).write_bytes(PNG)


def test_update_creates_all_outputs(tmp_path: Path) -> None:
    _write_project(tmp_path)
    result = Chapter6DissertationUpdater(tmp_path).build_and_save()
    assert result.document_path.exists()
    assert result.markdown_path.exists()
    assert result.report_json_path.exists()
    assert result.report_markdown_path.exists()


def test_update_creates_backup_for_in_place_edit(tmp_path: Path) -> None:
    _write_project(tmp_path)
    result = Chapter6DissertationUpdater(tmp_path).build_and_save()
    assert result.backup_path is not None and result.backup_path.exists()


def test_output_contains_eight_figures(tmp_path: Path) -> None:
    _write_project(tmp_path)
    Chapter6DissertationUpdater(tmp_path).build_and_save()
    assert len(Document(tmp_path / DEFAULT_DOCUMENT_NAME).inline_shapes) == 8


def test_output_has_no_placeholders(tmp_path: Path) -> None:
    _write_project(tmp_path)
    Chapter6DissertationUpdater(tmp_path).build_and_save()
    document = Document(tmp_path / DEFAULT_DOCUMENT_NAME)
    text = "\n".join(p.text for p in document.paragraphs)
    text += "\n" + "\n".join(c.text for t in document.tables for r in t.rows for c in r.cells)
    assert "[рассчитать]" not in text
    assert "[автозаполнение]" not in text
    assert "[n]" not in text


def test_report_records_stage_and_status(tmp_path: Path) -> None:
    _write_project(tmp_path)
    report = Chapter6DissertationUpdater(tmp_path).build_and_save().report
    assert report["stage"] == 15
    assert report["passed"] is True
    assert report["hypothesis_status"] == "hypothesis_partially_supported"


def test_report_records_nine_tables_and_eight_figures(tmp_path: Path) -> None:
    _write_project(tmp_path)
    report = Chapter6DissertationUpdater(tmp_path).build_and_save().report
    assert report["filled_tables"] == 9
    assert report["inserted_figures"] == 8


def test_markdown_contains_integral_metrics(tmp_path: Path) -> None:
    _write_project(tmp_path)
    result = Chapter6DissertationUpdater(tmp_path).build_and_save()
    text = result.markdown_path.read_text(encoding="utf-8")
    assert "0,159244" in text
    assert "hypothesis_partially_supported" in text


def test_rejected_acceptance_blocks_update(tmp_path: Path) -> None:
    _write_project(tmp_path)
    path = tmp_path / "reports/chapter6/chapter6_acceptance_report.json"
    payload = json.loads(path.read_text(encoding="utf-8"))
    payload["accepted"] = False
    path.write_text(json.dumps(payload), encoding="utf-8")
    with pytest.raises(Chapter6DissertationUpdateError, match="этапа 14"):
        Chapter6DissertationUpdater(tmp_path).build_and_save()


def test_wrong_row_count_blocks_update(tmp_path: Path) -> None:
    _write_project(tmp_path)
    path = tmp_path / "reports/chapter6/chapter6_acceptance_report.json"
    payload = json.loads(path.read_text(encoding="utf-8"))
    payload["row_count"] = 149
    path.write_text(json.dumps(payload), encoding="utf-8")
    with pytest.raises(Chapter6DissertationUpdateError, match="150"):
        Chapter6DissertationUpdater(tmp_path).build_and_save()


def test_missing_json_blocks_update(tmp_path: Path) -> None:
    _write_project(tmp_path)
    (tmp_path / "reports/chapter6/validation_metrics.json").unlink()
    with pytest.raises(FileNotFoundError, match="JSON"):
        Chapter6DissertationUpdater(tmp_path).build_and_save()


def test_missing_csv_blocks_update(tmp_path: Path) -> None:
    _write_project(tmp_path)
    (tmp_path / "reports/chapter6/baseline_comparison.csv").unlink()
    with pytest.raises(FileNotFoundError, match="CSV"):
        Chapter6DissertationUpdater(tmp_path).build_and_save()


def test_missing_figure_blocks_update(tmp_path: Path) -> None:
    _write_project(tmp_path)
    (tmp_path / "reports/chapter6/figures/confusion_matrix.png").unlink()
    with pytest.raises(FileNotFoundError, match="рисунок"):
        Chapter6DissertationUpdater(tmp_path).build_and_save()


def test_bad_figure_manifest_blocks_update(tmp_path: Path) -> None:
    _write_project(tmp_path)
    path = tmp_path / "reports/chapter6/figures/figure_manifest.json"
    path.write_text(json.dumps({"passed": False, "figure_count": 7}), encoding="utf-8")
    with pytest.raises(Chapter6DissertationUpdateError, match="восемь"):
        Chapter6DissertationUpdater(tmp_path).build_and_save()


def test_custom_output_does_not_replace_source(tmp_path: Path) -> None:
    _write_project(tmp_path)
    source = tmp_path / DEFAULT_DOCUMENT_NAME
    source_bytes = source.read_bytes()
    output = Path("reports/chapter6/dissertation/chapter6_updated.docx")
    result = Chapter6DissertationUpdater(tmp_path, output_document=output).build_and_save()
    assert source.read_bytes() == source_bytes
    assert result.backup_path is None
    assert result.document_path.exists()


def test_integral_table_is_filled(tmp_path: Path) -> None:
    _write_project(tmp_path)
    Chapter6DissertationUpdater(tmp_path).build_and_save()
    document = Document(tmp_path / DEFAULT_DOCUMENT_NAME)
    table = next(t for t in document.tables if t.rows[0].cells[0].text == "Метрика")
    assert table.rows[1].cells[1].text == "0,159244"


def test_confusion_matrix_is_filled(tmp_path: Path) -> None:
    _write_project(tmp_path)
    Chapter6DissertationUpdater(tmp_path).build_and_save()
    document = Document(tmp_path / DEFAULT_DOCUMENT_NAME)
    table = next(t for t in document.tables if t.rows[0].cells[0].text == "Фактический / прогнозный")
    assert [cell.text for cell in table.rows[2].cells[1:]] == ["66", "33", "1"]


def test_top_errors_are_filled(tmp_path: Path) -> None:
    _write_project(tmp_path)
    Chapter6DissertationUpdater(tmp_path).build_and_save()
    document = Document(tmp_path / DEFAULT_DOCUMENT_NAME)
    table = next(t for t in document.tables if t.rows[0].cells[0].text == "scenario_id")
    assert table.rows[1].cells[0].text == "S001"


def test_conclusion_contains_partial_support(tmp_path: Path) -> None:
    _write_project(tmp_path)
    Chapter6DissertationUpdater(tmp_path).build_and_save()
    text = "\n".join(p.text for p in Document(tmp_path / DEFAULT_DOCUMENT_NAME).paragraphs)
    assert "Основная гипотеза исследования подтверждена частично" in text


def test_source_reports_remain_unchanged(tmp_path: Path) -> None:
    _write_project(tmp_path)
    path = tmp_path / "reports/chapter6/validation_metrics.json"
    before = path.read_bytes()
    result = Chapter6DissertationUpdater(tmp_path).build_and_save()
    assert path.read_bytes() == before
    assert result.report["source_artifacts_modified"] is False


def test_second_run_is_reproducible(tmp_path: Path) -> None:
    _write_project(tmp_path)
    updater = Chapter6DissertationUpdater(tmp_path)
    updater.build_and_save()
    updater.build_and_save()
    document = Document(tmp_path / DEFAULT_DOCUMENT_NAME)
    assert len(document.inline_shapes) == 8


def test_cli_updates_document(
    tmp_path: Path,
    capsys: pytest.CaptureFixture[str],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _write_project(tmp_path)
    (tmp_path / "configs").mkdir()
    (tmp_path / "configs/chapter6.yaml").write_text(
        """
outputs:
  reports_dir: reports/chapter6
merge:
  key_columns: [scenario_id, protocol_id]
  validation: one_to_one
  expected_row_count: 150
decision_thresholds:
  low_max: 0.45
  high_min: 0.70
bootstrap:
  resamples: 1000
  confidence_level: 0.95
  random_seed: 42
  sampling_unit: scenario_id
""".strip(), encoding="utf-8")
    import manual_coding_sim.validation.chapter6_runner as runner

    monkeypatch.setattr(runner, "_load_config", lambda *_args, **_kwargs: object())
    monkeypatch.setattr(runner, "_print_config", lambda *_args, **_kwargs: None)
    code = main(["--project-root", str(tmp_path), "--update-dissertation"])
    assert code == 0
    assert "Этап 15 выполнен" in capsys.readouterr().out


def test_cli_returns_error_without_acceptance(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _write_project(tmp_path)
    (tmp_path / "reports/chapter6/chapter6_acceptance_report.json").unlink()
    (tmp_path / "configs").mkdir()
    (tmp_path / "configs/chapter6.yaml").write_text(
        """
outputs:
  reports_dir: reports/chapter6
merge:
  key_columns: [scenario_id, protocol_id]
  validation: one_to_one
  expected_row_count: 150
decision_thresholds:
  low_max: 0.45
  high_min: 0.70
bootstrap:
  resamples: 1000
  confidence_level: 0.95
  random_seed: 42
  sampling_unit: scenario_id
""".strip(), encoding="utf-8")
    import manual_coding_sim.validation.chapter6_runner as runner

    monkeypatch.setattr(runner, "_load_config", lambda *_args, **_kwargs: object())
    monkeypatch.setattr(runner, "_print_config", lambda *_args, **_kwargs: None)
    assert main(["--project-root", str(tmp_path), "--update-dissertation"]) == 1
